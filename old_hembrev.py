import datetime   
from docx import Document
from docx.shared import Inches
import os
from send import *
dt = datetime.date.today()
wk = dt.isocalendar()[1]
document = Document()





document.add_heading('Hembrev', 0)

p = document.add_paragraph('Skolarbete')

armatte = input("vad gör du i matten? ")
lermatte = input("Vad har du lärt dig i matten? ")
notematte = input("Har du någon note för matten? ")

arsv = input("Vad gör du i Svenskan? ")
lersv = input("Vad har du lärt dig i svenskan? ")
notesv = input("har du någon note för svenskan? ")

areng = input("Vad gör du i Engelskan? ")
lereng = input("Vad har du lärt dig i engelskan? ")
noteeng = input("har du någon note för engelskan? ")

arno = input("Vad gör du i NO? ")
lerno = input("Vad har du lärt dig i NO? ")
noteno = input("har du någon note för NO? ")

arso = input("Vad gör du i SO? ")
lerso = input("Vad har du lärt dig i SO? ")
noteso = input("har du någon note för SO? ")

aridh = input("Vad gör du i idrott? ")
leridh = input("Vad har du lärt dig i Idrotten? ")
noteidh = input("har du någon note för idrotten? ")

armu = input("Vad gör du i musiken? ")
lermu = input("Vad har du lärt dig i musiken? ")
notemu = input("har du någon note för Musik? ")

arbi = input("Vad gör du i Bilden? ")
lerbi = input("Vad har du lärt dig i Bilden? ")
notebi = input("har du någon note för Bild? ")

arhkk = input("Vad gör du i hemkunskapen? ")
lerhkk = input("Vad har du lärt dig i hemkunskapen? ")
notehkk = input("har du någon note för hemkunskap? ")

arsp = input("Vad gör du i spanskan? ")
lersp = input("Vad har du lärt dig i spanskan? ")
notesp = input("har du någon note för Spanskan? ")

arsl = input("Vad gör du i slöjden? ")
lersl = input("Vad har du lärt dig i slöjden? ")
notesl = input("har du någon note för slöjden? ")

rester = input("har du några rester (ja eller nej)")
if rester == "ja":
    print("rest 1/3")
    vad1 = input("Vad för rester har du? ")
    när1 = input("När ska du göra dem? ")
    hur1 = input("Hur ska du göra dem? ")
    print("rest 2/3")
    vad2 = input("vad för rester har du? ")
    när2 = input("När ska du göra dem? ")
    hur2 = input("Hur ska du göra dem? ")
    print("rest 3/3")
    vad3 = input("Vad för rester har du? ")
    när3 = input("När ska du göra dem? ")
    hur3 = input("Hur ska du göra dem? ")
else:
    vad1 = " "
    hur1 = " "
    när1 = " "
    vad2 = " "
    hur2 = " "
    när2 = " "
    vad3 = " "
    hur3 = " "
    när3 = " "
    
    
    
hend = input("Har ni något som kommer hända (ja eller nej)? ")
if hend == "ja":
    print("händelse 1/3")
    hendvad1 = input("Vad ska ni göra? ")
    hendnär1 = input("När ska ni göra det? ")
    print("händelse 2/3")
    hendvad2 = input("Vad ska ni göra? ")
    hendnär2 = input("När ska ni göra det? ")
    print("händelse 3/3")
    hendvad3 = input("Vad ska ni göra? ")
    hendnär3 = input("När ska ni göra det? ")
else:
    hendvad1 = " "
    hendnär1 = " "
    hendvad2 = " "
    hendnär2 = " "
    hendvad3 = " "
    hendnär3 = " "
if notematte == "nej" or notematte == "no":
    notematte = "-"
if notesv == "nej" or notesv == "no":
    notesv = "-"
if noteeng == "nej" or noteeng == "no":
    noteeng = "-"
if noteno == "nej" or noteno == "no":
    noteno = "-"
if noteso == "nej" or noteso == "no":
    noteso = "-"
if noteidh == "nej" or noteidh == "no":
    noteidh = "-"
if notemu == "nej" or notemu == "no":
    notemu = "-"
if noteidh == "nej" or noteidh == "no":
    noteidh = "-"
if notebi == "nej" or notebi == "no":
    notebi = "-"
if notehkk == "nej" or notehkk == "no":
    notehkk = "-"
if notesl == "nej" or notesl == "no":
    notesl = "-"
if notesp == "nej" or notesp == "no":
    notesp = "-"



records = (
    ("Matte", armatte, lermatte, notematte),
    ("svenska", arsv, lersv, notesv),
    ("Engelska", areng, lereng, noteeng), 
    ("NO", arno, lerno, noteno),
    ("SO", arso, lerso, noteso),
    ("Idrott", aridh, leridh, noteidh),
    ("Musik", armu, lermu, notemu),
    ("Bild", arbi, lerbi, notebi),
    ("Hemkundskap", arhkk, lerhkk, notehkk),
    ("Slöjd", arsl, lersl, notesl),
    ("Spanska", arsp, lersp, notesp),
    
)
rester = (
    ( vad1, när1, hur1),
    ( vad2, när2, hur2),
    ( vad3, när3, hur3)
    )
hend = (
    ( hendvad1, hendnär1),
    ( hendvad2, hendnär2),
    ( hendvad3, hendnär3)
)

table1 = document.add_table(rows=1, cols=4)
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Ämne:'
hdr_cells[1].text = 'Arbetsområde:'
hdr_cells[2].text = 'Detta har jag lärt mig:'
hdr_cells[3].text = 'Notes:'
for qty, id, desc, des in records:
    row_cells = table1.add_row().cells
    row_cells[0].text = qty
    row_cells[1].text = id
    row_cells[2].text = desc
    row_cells[3].text = des
document.add_heading('Rester', 0)
table2 = document.add_table(rows=1, cols=3)
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Rester'
hdr_cells[1].text = 'När'
hdr_cells[2].text = 'Hur'
for vad1, vad2, vad3 in rester:
    row_cells = table2.add_row().cells
    row_cells[0].text = vad1
    row_cells[1].text = vad2
    row_cells[2].text = vad3
document.add_heading('Händelser', 0)
table3 = document.add_table(rows=1, cols=2)
hdr_cells = table3.rows[0].cells
hdr_cells[0].text = 'vad'
hdr_cells[1].text = 'När'
for vad1, vad2, vad3 in rester:
    row_cells = table3.add_row().cells
    row_cells[0].text = hendvad1
    row_cells[1].text = hendvad2
document.add_page_break()

document.save('hembrev v' + str(wk) + ".docx")


sendmail()





