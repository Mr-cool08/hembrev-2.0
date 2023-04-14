import datetime   
from docx import Document
from docx.shared import Inches
import os
#from send import *
dt = datetime.date.today()
wk = dt.isocalendar()[1]
document = Document()





document.add_heading('Hembrev', 0)

p = document.add_paragraph('Skolarbete')

arMatte = input("vad gör du i Matten? ")
lerMatte = input("Vad har du lärt dig i Matten? ")
noteMatte = input("Har du någon note för Matten? ")

arsv = input("Vad gör du i Svenskan? ")
lersv = input("Vad har du lärt dig i svenskan? ")
notesv = input("har du någon note för svenskan? ")

areng = input("Vad gör du i Engelskan? ")
lereng = input("Vad har du lärt dig i engelskan? ")
noteeng = input("har du någon note för engelskan? ")

arno = input("Vad gör du i NO? ")
lerno = input("Vad har du lärt dig i NO? ")
noteno = input("har du någon note för NO? ")

arso = input("Vad gör du i SO? ")
lerso = input("Vad har du lärt dig i SO? ")
noteso = input("har du någon note för SO? ")

aridh = input("Vad gör du i idrott? ")
leridh = input("Vad har du lärt dig i Idrotten? ")
noteidh = input("har du någon note för idrotten? ")

armu = input("Vad gör du i musiken? ")
lermu = input("Vad har du lärt dig i musiken? ")
notemu = input("har du någon note för Musik? ")

arbi = input("Vad gör du i Bilden? ")
lerbi = input("Vad har du lärt dig i Bilden? ")
notebi = input("har du någon note för Bild? ")

arhkk = input("Vad gör du i hemkunskapen? ")
lerhkk = input("Vad har du lärt dig i hemkunskapen? ")
notehkk = input("har du någon note för hemkunskap? ")

arsp = input("Vad gör du i spanskan? ")
lersp = input("Vad har du lärt dig i spanskan? ")
notesp = input("har du någon note för Spanskan? ")

arsl = input("Vad gör du i slöjden? ")
lersl = input("Vad har du lärt dig i slöjden? ")
notesl = input("har du någon note för slöjden? ")

rester = input("har du några rester (ja eller nej)")
if rester == "ja":
    print("rest 1/3")
    vad1 = input("Vad för rester har du? ")
    när1 = input("När ska du göra dem? ")
    hur1 = input("Hur ska du göra dem? ")
    print("rest 2/3")
    vad2 = input("vad för rester har du? ")
    när2 = input("När ska du göra dem? ")
    hur2 = input("Hur ska du göra dem? ")
    print("rest 3/3")
    vad3 = input("Vad för rester har du? ")
    när3 = input("När ska du göra dem? ")
    hur3 = input("Hur ska du göra dem? ")
else:
    vad1 = " "
    hur1 = " "
    när1 = " "
    vad2 = " "
    hur2 = " "
    när2 = " "
    vad3 = " "
    hur3 = " "
    när3 = " "
    
    
    
hend = input("Har ni något som kommer hända (ja eller nej)? ")
if hend == "ja":
    print("händelse 1/3")
    hendvad1 = input("Vad ska ni göra? ")
    hendnär1 = input("När ska ni göra det? ")
    print("händelse 2/3")
    hendvad2 = input("Vad ska ni göra? ")
    hendnär2 = input("När ska ni göra det? ")
    print("händelse 3/3")
    hendvad3 = input("Vad ska ni göra? ")
    hendnär3 = input("När ska ni göra det? ")
else:
    hendvad1 = " "
    hendnär1 = " "
    hendvad2 = " "
    hendnär2 = " "
    hendvad3 = " "
    hendnär3 = " "
if noteMatte == "nej" or noteMatte == "no":
    noteMatte = "-"
if notesv == "nej" or notesv == "no":
    notesv = "-"
if noteeng == "nej" or noteeng == "no":
    noteeng = "-"
if noteno == "nej" or noteno == "no":
    noteno = "-"
if noteso == "nej" or noteso == "no":
    noteso = "-"
if noteidh == "nej" or noteidh == "no":
    noteidh = "-"
if notemu == "nej" or notemu == "no":
    notemu = "-"
if noteidh == "nej" or noteidh == "no":
    noteidh = "-"
if notebi == "nej" or notebi == "no":
    notebi = "-"
if notehkk == "nej" or notehkk == "no":
    notehkk = "-"
if notesl == "nej" or notesl == "no":
    notesl = "-"
if notesp == "nej" or notesp == "no":
    notesp = "-"



records = (
    ("Matte", arMatte, lerMatte, noteMatte),
    ("svenska", arsv, lersv, notesv),
    ("Engelska", areng, lereng, noteeng), 
    ("NO", arno, lerno, noteno),
    ("SO", arso, lerso, noteso),
    ("Idrott", aridh, leridh, noteidh),
    ("Musik", armu, lermu, notemu),
    ("Bild", arbi, lerbi, notebi),
    ("Hemkundskap", arhkk, lerhkk, notehkk),
    ("Slöjd", arsl, lersl, notesl),
    ("Spanska", arsp, lersp, notesp),
    
)
rester = (
    ( vad1, när1, hur1),
    ( vad2, när2, hur2),
    ( vad3, när3, hur3)
    )
hend = (
    ( hendvad1, hendnär1),
    ( hendvad2, hendnär2),
    ( hendvad3, hendnär3)
)

table1 = document.add_table(rows=1, cols=4)
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Ämne:'
hdr_cells[1].text = 'Arbetsområde:'
hdr_cells[2].text = 'Detta har jag lärt mig:'
hdr_cells[3].text = 'Notes:'
for qty, id, desc, des in records:
    row_cells = table1.add_row().cells
    row_cells[0].text = qty
    row_cells[1].text = id
    row_cells[2].text = desc
    row_cells[3].text = des
document.add_heading('Rester', 0)
table2 = document.add_table(rows=1, cols=3)
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Rester'
hdr_cells[1].text = 'När'
hdr_cells[2].text = 'Hur'
for vad1, vad2, vad3 in rester:
    row_cells = table2.add_row().cells
    row_cells[0].text = vad1
    row_cells[1].text = vad2
    row_cells[2].text = vad3
document.add_heading('Händelser', 0)
table3 = document.add_table(rows=1, cols=2)
hdr_cells = table3.rows[0].cells
hdr_cells[0].text = 'vad'
hdr_cells[1].text = 'När'
for vad1, vad2, vad3 in rester:
    row_cells = table3.add_row().cells
    row_cells[0].text = hendvad1
    row_cells[1].text = hendvad2
document.add_page_break()

document.save('hembrev v' + str(wk) + ".docx")


#sendmail()


data = [{ "tag": "Matte", "patterns": ["vad g\u00f6r du i Matten"], "responses": [arMatte], "context_set": "" },
{ "tag": "Matte", "patterns": ["Vad har du lärt dig"], "responses": [lerMatte], "context_set": "" },
{ "tag": "Matte", "patterns": ["Har du någon note för Matten?"], "responses": [noteMatte], "context_set": "" },

{ "tag": "svenska", "patterns": ["vad g\u00f6r du i Svenskan"], "responses": [arsv], "context_set": "" },
{ "tag": "svenska", "patterns": ["Vad har du lärt dig i Svenskan"], "responses": [lersv], "context_set": "" },
{ "tag": "svenska", "patterns": ["Har du någon note för Svenskan?"], "responses": [notesv], "context_set": "" },

{ "tag": "eng", "patterns": ["vad g\u00f6r du i Engelskan"], "responses": [areng], "context_set": "" },
{ "tag": "eng", "patterns": ["Vad har du lärt dig i Engelskan?"], "responses": [lereng], "context_set": "" },
{ "tag": "eng", "patterns": ["Har du någon note för Engelskan?"], "responses": [noteeng], "context_set": "" },

{ "tag": "no", "patterns": ["vad g\u00f6r du i NO"], "responses": [arno], "context_set": "" },
{ "tag": "no", "patterns": ["Vad har du lärt dig i NO"], "responses": [lerno], "context_set": "" },
{ "tag": "no", "patterns": ["Har du någon note för NO?"], "responses": [noteno], "context_set": "" },

{ "tag": "SO", "patterns": ["vad g\u00f6r du i SO"], "responses": [arso], "context_set": "" },
{ "tag": "SO", "patterns": ["Vad har du lärt dig i SO"], "responses": [lerso], "context_set": "" },
{ "tag": "SO", "patterns": ["Har du någon note för SO?"], "responses": [noteso], "context_set": "" },

{ "tag": "idh", "patterns": ["vad g\u00f6r du i Idrotten"], "responses": [aridh], "context_set": "" },
{ "tag": "idh", "patterns": ["Vad har du lärt dig i Idrotten"], "responses": [leridh], "context_set": "" },
{ "tag": "idh", "patterns": ["Har du någon note för Idrotten?"], "responses": [noteidh], "context_set": "" },

{ "tag": "mu", "patterns": ["vad g\u00f6r du i Musiken"], "responses": [armu], "context_set": "" },
{ "tag": "mu", "patterns": ["Vad har du lärt dig i Musiken"], "responses": [lermu], "context_set": "" },
{ "tag": "mu", "patterns": ["Har du någon note för Musiken?"], "responses": [notemu], "context_set": "" },

{ "tag": "bi", "patterns": ["vad g\u00f6r du i Bilden"], "responses": [arbi], "context_set": "" },
{ "tag": "bi", "patterns": ["Vad har du lärt dig i Bilden"], "responses": [lerbi], "context_set": "" },
{ "tag": "bi", "patterns": ["Har du någon note för Bilden?"], "responses": [notebi], "context_set": "" },

{ "tag": "hkk", "patterns": ["vad g\u00f6r du i hemkunskapen"], "responses": [arhkk], "context_set": "" },
{ "tag": "hkk", "patterns": ["Vad har du lärt dig i hemkunskapen"], "responses": [lerhkk], "context_set": "" },
{ "tag": "hkk", "patterns": ["Har du någon note för hemkunskapen?"], "responses": [notehkk], "context_set": "" },

{ "tag": "sl", "patterns": ["vad g\u00f6r du i Slöjden"], "responses": [arsl], "context_set": "" },
{ "tag": "sl", "patterns": ["Vad har du lärt dig i Slöjden"], "responses": [lersl], "context_set": "" },
{ "tag": "sl", "patterns": ["Har du någon note för Slöjden?"], "responses": [notesl], "context_set": "" },

{ "tag": "sp", "patterns": ["vad g\u00f6r du i Spanskan"], "responses": [arsp], "context_set": "" },
{ "tag": "sp", "patterns": ["Vad har du lärt dig i Spanskan"], "responses": [lersp], "context_set": "" },
{ "tag": "sp", "patterns": ["Har du någon note för Spanskan?"], "responses": [notesp], "context_set": "" },
]
 
import json
 
 
# function to add to JSON
 
def write_json(new_data, filename='hembrev.json'):
    with open(filename,'r+') as file:
          # First we load existing data into a dict.
        file_data = json.load(file)
        # Join new_data with file_data inside emp_details
        file_data["intents"].append(new_data)
        # Sets file's current position at offset.
        file.seek(0)
        # convert back to json.
        json.dump(file_data, file, indent = 4)
 
    # python object to be appended

     
write_json(data)