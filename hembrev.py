from tkinter import ttk
from docx import Document
import datetime
import socket
import smtplib
import logging
import glob
from tkinter import messagebox
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from windtalker import SymmetricCipher
import configparser
import random
import tkinter as tk
from configparser import ConfigParser
import time
import requests
import urllib.request
import os
import socket
host_name = socket.gethostname()
#getting the week number
dt = datetime.date.today()
wk = dt.isocalendar()[1]
path = os.path.join(os.path.expanduser('~'), 'AppData', 'Roaming')
newPath = path.replace(os.sep, '/')
fullpath = newPath + "/hembrev"
logfile = f"{fullpath}/ErrorLog {host_name} {dt}.log"
logging.basicConfig(filename=logfile, level=logging.ERROR,
                    format='%(asctime)s - %(levelname)s - %(message)s')
def upload_file():
    """
    Uploads the "error.log" file to a Flask app via a POST request.

    Returns:
        str: The response from the Flask app.
    """
    file_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), f'Error {host_name} {dt}.log')
    # Make sure the file exists before attempting to upload
    if os.path.exists(file_path):
        url = 'https://hembrev-3.mrcoolcool.repl.co/upload'
        files = {'file': open(file_path, 'rb')} # Open the file in binary mode

        response = requests.post(url, files=files)

        if response.status_code == 200:
            return 'File uploaded successfully'
        else:
            return 'File upload failed'
    else:
        return 'File not found'


def decrypt():
    # Read the password from file
    with open(fullpath + "/file.txt", "r") as file:
        password = file.read().strip()

    c = SymmetricCipher(password=password)
    try:
        folder_path = os.path.join(newPath, "hembrev")  # Combine newPath and folder name
        filename = os.path.join(folder_path, "password-encrypted.txt")
        c.decrypt_file(filename)
        os.remove(filename)
    except OSError as e:
        print(e)
        folder_path = os.path.join(newPath, "hembrev").replace("\\", "/")  # Combine newPath and folder name and replace backslashes with forward slashes
        filename = os.path.join(folder_path, "password.txt").replace("\\", "/")
        c.encrypt_file(filename)
        os.remove(filename)
        decrypt()

class Application(tk.Frame):
    def __init__(self, master=None):
        super().__init__(master)
        self.master = master
        self.grid()
        self.create_widgets()
    
      
    #mark these as empty if the user dont input    
    händelse_vad1 = ""
    händelse_när1 = ""
    händelse_vad2 = ""
    händelse_när2 = ""
    händelse_vad3 = ""
    händelse_när3 = ""
    rest_vad1 = ""
    rest_hur1 = ""
    rest_när1 = ""
    rest_vad2 = ""
    rest_hur2 = ""
    rest_när2 = ""
    rest_vad3 = ""
    rest_hur3 = ""
    rest_när3 = ""
   
    
    
    
    
    #encrypt and decrypt for the password
    def encrypt(self):
    # Read the password from file
        with open(fullpath + "/file.txt", "r") as file:
            password = file.read().strip()

        c = SymmetricCipher(password=password)
        try:
            folder_path = os.path.join(newPath, "hembrev")  # Combine newPath and folder name
            filename = os.path.join(folder_path, "password.txt")
            c.encrypt_file(filename)
            os.remove(filename)
        except OSError as e:
            print(e)
            os.remove("password-encrypted.txt")
            self.encrypt()
    def decrypt(self):
        # Read the password from file
        with open(fullpath + "/file.txt", "r") as file:
            password = file.read().strip()

        c = SymmetricCipher(password=password)
        try:
            folder_path = os.path.join(newPath, "hembrev")  # Combine newPath and folder name
            filename = os.path.join(folder_path, "password-encrypted.txt")
            c.decrypt_file(filename)
            os.remove(filename)
        except OSError as e:
            print(e)
            folder_path = os.path.join(newPath, "hembrev").replace("\\", "/")  # Combine newPath and folder name and replace backslashes with forward slashes
            filename = os.path.join(folder_path, "password.txt").replace("\\", "/")
            c.encrypt_file(filename)
            os.remove(filename)
            self.decrypt()
            
    

    
    #if the user want to send something in the mail
    def mail_message(self):
        self.title_label = tk.Label(self, text="Skriv här vad som ska stå i mailet")
        self.title_label.grid(row=23, column=1, pady=10, columnspan=2)
        frame = tk.Frame(self)
        frame.grid(row=24, column=1, sticky="w", columnspan=50)
        self.message = tk.Entry(frame, width=50)
        self.message.grid(row=0, column=0)
        self.message.pack()



        



    
    #if the user have any rester
    def rest(self):
            
            
        self.title_label = tk.Label(self, text="Vad")
        self.title_label.grid(row=14, column=0, columnspan=3, pady=10)
        self.title_label = tk.Label(self, text="Planering")
        self.title_label.grid(row=14, column=1, columnspan=5, pady=10)
        self.title_label = tk.Label(self, text="När")
        self.title_label.grid(row=14, column=3, columnspan=6, pady=10)
        self.extra_label1 = tk.Label(self, text="Rest 1")
        self.extra_label1.grid(row=15, column=0)

        self.rest_vad1 = tk.Entry(self)
        self.rest_vad1.grid(row=15, column=1)
        self.rest_vad1.bind("<Return>", lambda event: self.rest_hur1.focus_set())

        self.rest_hur1 = tk.Entry(self)
        self.rest_hur1.grid(row=15, column=2)
        self.rest_hur1.bind("<Return>", lambda event: self.rest_när1.focus_set())

        self.rest_när1 = tk.Entry(self)
        self.rest_när1.grid(row=15, column=3)
        self.rest_när1.bind("<Return>", lambda event: self.rest_vad2.focus_set())
        
        
        self.extra_label2 = tk.Label(self, text="Rest 2")
        self.extra_label2.grid(row=16, column=0)

        self.rest_vad2 = tk.Entry(self)
        self.rest_vad2.grid(row=16, column=1)
        self.rest_vad2.bind("<Return>", lambda event: self.rest_hur2.focus_set())

        self.rest_hur2 = tk.Entry(self)
        self.rest_hur2.grid(row=16, column=2)
        self.rest_hur2.bind("<Return>", lambda event: self.rest_när2.focus_set())
        
        self.rest_när2 = tk.Entry(self)
        self.rest_när2.grid(row=16, column=3)
        self.rest_när2.bind("<Return>", lambda event: self.rest_vad3.focus_set())
        
        
        self.extra_label3 = tk.Label(self, text="Rest 3")
        self.extra_label3.grid(row=17, column=0)

        self.rest_vad3 = tk.Entry(self)
        self.rest_vad3.grid(row=17, column=1)
        self.rest_vad3.bind("<Return>", lambda event: self.rest_hur3.focus_set())

        self.rest_hur3 = tk.Entry(self)
        self.rest_hur3.grid(row=17, column=2)
        self.rest_hur3.bind("<Return>", lambda event: self.rest_när3.focus_set())

        self.rest_när3 = tk.Entry(self)
        self.rest_när3.grid(row=17, column=3)
        tk.Entry.pack(self)
        
        
       
    #if the user have anything that is going to happen    
    def händelser(self):
            
            
        self.title_label = tk.Label(self, text="Vad")
        self.title_label.grid(row=18, column=0, columnspan=3, pady=10)
        self.title_label = tk.Label(self, text="När")
        self.title_label.grid(row=18, column=1, columnspan=3, pady=10)
        
        self.extra_label1 = tk.Label(self, text="Händelse 1")
        self.extra_label1.grid(row=19, column=0)

        self.händelse_vad1 = tk.Entry(self)
        self.händelse_vad1.grid(row=19, column=1)
        self.händelse_vad1.bind("<Return>", lambda event: self.händelse_när1.focus_set())

        self.händelse_när1 = tk.Entry(self)
        self.händelse_när1.grid(row=19, column=2)
        self.händelse_när1.bind("<Return>", lambda event: self.händelse_vad2.focus_set())

        
        self.extra_label2 = tk.Label(self, text="Händelse 2")
        self.extra_label2.grid(row=20, column=0)

        self.händelse_vad2 = tk.Entry(self)
        self.händelse_vad2.grid(row=20, column=1)
        self.händelse_vad2.bind("<Return>", lambda event: self.händelse_när2.focus_set())

        self.händelse_när2 = tk.Entry(self)
        self.händelse_när2.grid(row=20, column=2)
        self.händelse_när2.bind("<Return>", lambda event: self.händelse_vad3.focus_set())
        
        
        self.extra_label3 = tk.Label(self, text="Händelse 3")
        self.extra_label3.grid(row=21, column=0)

        self.händelse_vad3 = tk.Entry(self)
        self.händelse_vad3.grid(row=21, column=1)
        self.händelse_vad3.bind("<Return>", lambda event: self.händelse_när3.focus_set())

        self.händelse_när3 = tk.Entry(self)
        self.händelse_när3.grid(row=21, column=2)
        
        
        
     #creating the document with all inputs
    def create_document(self):
        document = Document()

        document.add_heading('Hembrev', 0)

        # Create table for all subjects
        subject_data = [
            ('Matte', self.matte_gör.get(), self.matte_lärt.get(), self.matte_note.get()),
            ('Svenska', self.sv_gör.get(), self.sv_lärt.get(), self.sv_note.get()),
            ('Engelska', self.eng_gör.get(), self.eng_lärt.get(), self.eng_note.get()),
            ('NO', self.no_gör.get(), self.no_lärt.get(), self.no_note.get()),
            ('SO', self.so_gör.get(), self.so_lärt.get(), self.so_note.get()),
            ('Idrott', self.idh_gör.get(), self.idh_lärt.get(), self.idh_note.get()),
            ('Musik', self.mu_gör.get(), self.mu_lärt.get(), self.mu_note.get()),
            ('Bild', self.bi_gör.get(), self.bi_lärt.get(), self.bi_note.get()),
            ('Hemkunskap', self.hkk_gör.get(), self.hkk_lärt.get(), self.hkk_note.get()),
            ('Språk', self.sp_gör.get(), self.sp_lärt.get(), self.sp_note.get()),
            ('Slöjd', self.sl_gör.get(), self.sl_lärt.get(), self.sl_note.get())
        ]

        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Ämne'
        hdr_cells[1].text = 'ARBETSOMRÅDE'
        hdr_cells[2].text = 'DETTA HAR JAG LÄRT MIG:'
        hdr_cells[3].text = 'notes'

        for subject_row in subject_data:
            row_cells = table.add_row().cells
            for i, value in enumerate(subject_row):
                row_cells[i].text = value

        # Create table for "händelser"
        document.add_heading('Händelser', 0)
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Vad'
        hdr_cells[1].text = 'När'

        if self.händelser_checked.get():
            händelse_data = [
                (self.händelse_vad1.get(), self.händelse_när1.get()),
                (self.händelse_vad2.get(), self.händelse_när2.get()),
                (self.händelse_vad3.get(), self.händelse_när3.get())
            ]

            for händelse_row in händelse_data:
                row_cells = table.add_row().cells
                for i, value in enumerate(händelse_row):
                    row_cells[i].text = value

        # Create table for "rester"
        document.add_heading('Rester', 0)
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Vad'
        hdr_cells[1].text = 'Planering'
        hdr_cells[2].text = 'När'

        if self.rest_checked.get():
            rest_data = [
                (self.rest_vad1.get(), self.rest_hur1.get(), self.rest_när1.get()),
                (self.rest_vad2.get(), self.rest_hur2.get(), self.rest_när2.get()),
                (self.rest_vad3.get(), self.rest_hur3.get(), self.rest_när3.get())
            ]

            for rest_row in rest_data:
                row_cells = table.add_row().cells
                for i, value in enumerate(rest_row):
                    row_cells[i].text = value

        path = os.path.expanduser('~')
        path1 = os.path.join(path, "hembrev")
        os.makedirs(path1, exist_ok=True)

        # Saving the document with the week number
        document.save(fr'{path1}\\Hembrev v{wk}.docx')

        
    
    def sendmail(self, message):
            #code for message in mail
            config = configparser.ConfigParser()
            config_file_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), f'{fullpath}/config.ini')
            print(config_file_path)
            config.read(config_file_path)
            mail = config.get('login', 'Email')
            name = mail.split(".")[0].capitalize()
            good_bye = ["Ha en bra dag!:)", f"Mvh {name}", f"vänliga hälsningar {name}", "Hoppas vi ses snart igen!", "Allt gott!", "Ta hand om dig!", "Vi hörs!", "Ha det så bra!", "Fortsatt trevlig dag!",
"Ha det fint!", "Kram!", "Adjö!", "På återseende!", "Tack för allt!", "Hoppas du får en bra dag!", "Ta hand om dig själv och varandra!", "Ha en fantastisk dag!", "Trevlig vecka!", "Hoppas vi ses snart igen!",
"Ha det bra på din resa!", "Lycka till!", "Vi hörs snart igen!", "Hoppas allt går bra för dig!", "Hoppas du får en bra dag!", "Det var trevligt att prata med er!", "Ta hand om dig själv!",
"Fortsätt ha en bra dag!", "Var rädd om dig!", "Ha en underbar dag!", "Hoppas du har en bra vecka!", "Ha det så trevligt!", "Puss och kram!", "Vi ses snart!",
"Hoppas du får en trevlig dag!", "Ha en fin dag!", "Hoppas allt går som du önskar!", "Tack för att du stannade upp en stund och läste hembrevet!", "Vi hörs igen snart!", "Hoppas du har en fantastisk dag!",
"Ha en skön dag!", "Ha det gott!", "Vi ses vid ett annat tillfälle!",
"Ta hand om dig själv och ha det bra!", "Hoppas du har en underbar dag!", "Ha en skön vecka!",
"Hoppas vi kan stödja dig igen snart!", "Hoppas du får en lyckad dag!", "Ha det så kul!",
"Ha en bra resa!", "Vi ses snart igen med nya möjligheter!"]
           
            
            self.decrypt()
            config = configparser.ConfigParser()
            config.read(f"{fullpath}/config.ini")
            mail1 = config.get('Emails', 'email1')
            mail2 = config.get('Emails', 'email2')
            mail3 = config.get('Emails', 'Email3')
            mail4 = config.get('Emails', 'Email4')
            mail5 = config.get('Emails', 'Email5')
            mail = config.get('login', 'Email')
            auto_message = config.get('Settings', 'auto_message')
            random_message = random.choice(good_bye)
            
            if auto_message == "0":
                random_message = " "
            
                
            
            mail_body = message  + "\n" + "\n" +"\n" +"\n" +"\n" + random_message +"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"Detta mail var skickat igenom hembrevs programmet"
            folder_path = os.path.join(newPath, "hembrev").replace("\\", "/")  # Combine newPath and folder name and replace backslashes with forward slashes
            filename = os.path.join(folder_path, "password.txt").replace("\\", "/")

            
            with open(filename,'r') as file:
                password = file.read()

            #making the email
            home_dir = os.path.expanduser("~")

            # Combine the home directory with the relative path to the folder
            folder_path = os.path.join(home_dir, "hembrev")

            # Replace backslashes with forward slashes
            folder_path = folder_path.replace("\\", "/")    
            docname = os.path.join(folder_path, f"hembrev v{wk}.docx").replace("\\", "/")
            msg = MIMEMultipart()
            msg['From'] = mail
            receivers = [mail1, mail2, mail3, mail4, mail5]
            msg['To'] = ', '.join(receivers)
            msg['Subject'] = f'hembrev v{wk}'
            body = mail_body
            msg.attach(MIMEText(body, 'plain'))
            with open(docname, 'rb') as attachment:
                part = MIMEBase('application', "octet-stream")
                part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f"attachment; filename=hembrev v{wk}.docx")
            msg.attach(part)


            #sending the mail with a office 365 server
            #sending the mail with a office 365 server
            try:
                server = smtplib.SMTP('smtp.office365.com', 587)  ### put your relevant SMTP here
            except (smtplib.SMTPSenderRefused, socket.gaierror) as e:
                messagebox.showerror("Error", f"Du har blivit bort kopplad från internet så kan inte komma åt serverar \n {e}")
                logging.error('An error occurred: %s', e)
                path = os.path.expanduser('~')
                path1 = path + "\\hembrev\\"
                path = os.path.realpath(path1)
                os.startfile(path)
                result = upload_file()
                print(result)
                self.encrypt()

            except Exception as e:
                messagebox.showerror("Error", f"Error kontakta Liam Suorsa \n {e}")
                logging.error('An error occurred: %s', e)
                path = os.path.expanduser('~')
                path1 = path + "\\hembrev\\"
                path = os.path.realpath(path1)
                os.startfile(path)
                result = upload_file()
                print(result)
                self.encrypt()
            server.ehlo()
            server.starttls()
            server.ehlo()
            try:
                server.login(mail, password) 
            except smtplib.SMTPAuthenticationError as e:
                messagebox.showerror("Error", f"Du har skrivit fel inloggnings uppgifter starta programmet igen och tryck på ändra epost \n {e}")
                logging.error('An error occurred: %s', e)
                path = os.path.expanduser('~')
                path1 = path + "\\hembrev\\"
                path = os.path.realpath(path1)
                os.startfile(path)
                result = upload_file()
                print(result)
                self.encrypt()
            except Exception as e:
                messagebox.showerror("Error", f"Error kontakta Liam Suorsa \n {e}")
                logging.error('An error occurred: %s', e)
                path = os.path.expanduser('~')
                path1 = path + "\\hembrev\\"
                path = os.path.realpath(path1)
                os.startfile(path)
                result = upload_file()
                print(result)
                self.encrypt()
            ### if applicable
            server.send_message(msg)
            server.quit()
            self.encrypt()
            print("done")
    #when the user presses the sumbit button
    def submit(self):
        try:
            if self.mail_message_checked.get() == 1:
                message = self.message.get().strip() or ' '
            else:
                message = ' '
            self.create_document()
            self.master.destroy()
            self.sendmail(message)
        except Exception as e:
            messagebox.showerror("Error", f"Error kontakta Liam Suorsa \n {e}")
            logging.error('An error occurred while submitting: %s', e)
            path = os.path.expanduser('~')
            path1 = path + "\\hembrev\\"
            path = os.path.realpath(path1)
            os.startfile(path)
            result = upload_file()
            print(result)
            



    
    #create_widgets is where all the visual stuff is
    def create_widgets(self):
        
        #labels so the user know where to input what
        self.title_label = tk.Label(self, text="Vad gör du?")
        self.title_label.grid(row=0, column=1, pady=10, sticky="w")
        self.title_label = tk.Label(self, text="Vad lär du dig?")
        self.title_label.grid(row=0, column=2, columnspan=6, pady=10, sticky="w")
        self.title_label = tk.Label(self, text="Note")
        self.title_label.grid(row=0, column=3, columnspan=7, pady=10, sticky="w")


        #the input boxes for the subjects
        def questions(self):
            self.matte_label = tk.Label(self, text="Matte")
            self.matte_label.grid(row=1, column=0, sticky="w")

            self.matte_gör = tk.Entry(self)
            self.matte_gör.grid(row=1, column=1, sticky="w")
            self.matte_gör.bind("<Return>", lambda event: self.matte_lärt.focus_set())

            self.matte_lärt = tk.Entry(self)
            self.matte_lärt.grid(row=1, column=2, sticky="w")
            self.matte_lärt.bind("<Return>", lambda event: self.matte_note.focus_set())

            self.matte_note = tk.Entry(self)
            self.matte_note.grid(row=1, column=3, sticky="w")
            self.matte_note.bind("<Return>", lambda event: self.sv_gör.focus_set())

            self.sv_label = tk.Label(self, text="Svenska")
            self.sv_label.grid(row=2, column=0, sticky="w")

            self.sv_gör = tk.Entry(self)
            self.sv_gör.grid(row=2, column=1, sticky="w")
            self.sv_gör.bind("<Return>", lambda event: self.sv_lärt.focus_set())

            self.sv_lärt = tk.Entry(self)
            self.sv_lärt.grid(row=2, column=2, sticky="w")
            self.sv_lärt.bind("<Return>", lambda event: self.sv_note.focus_set())

            self.sv_note = tk.Entry(self)
            self.sv_note.grid(row=2, column=3, sticky="w")
            self.sv_note.bind("<Return>", lambda event: self.eng_gör.focus_set())
            
            
            self.eng_label = tk.Label(self, text="Engelska")
            self.eng_label.grid(row=3, column=0, sticky="w")

            self.eng_gör = tk.Entry(self)
            self.eng_gör.grid(row=3, column=1)
            self.eng_gör.bind("<Return>", lambda event: self.eng_lärt.focus_set())

            self.eng_lärt = tk.Entry(self)
            self.eng_lärt.grid(row=3, column=2)
            self.eng_lärt.bind("<Return>", lambda event: self.eng_note.focus_set())

            self.eng_note = tk.Entry(self)
            self.eng_note.grid(row=3, column=3)
            self.eng_note.bind("<Return>", lambda event: self.no_gör.focus_set())
            
            self.no_label = tk.Label(self, text="NO")
            self.no_label.grid(row=4, column=0, sticky="w")

            self.no_gör = tk.Entry(self)
            self.no_gör.grid(row=4, column=1)
            self.no_gör.bind("<Return>", lambda event: self.no_lärt.focus_set())

            self.no_lärt = tk.Entry(self)
            self.no_lärt.grid(row=4, column=2)
            self.no_lärt.bind("<Return>", lambda event: self.no_note.focus_set())

            self.no_note = tk.Entry(self)
            self.no_note.grid(row=4, column=3)
            self.no_note.bind("<Return>", lambda event: self.so_gör.focus_set())
            
            
            self.so_label = tk.Label(self, text="SO")
            self.so_label.grid(row=5, column=0, sticky="w")

            self.so_gör = tk.Entry(self)
            self.so_gör.grid(row=5, column=1)
            self.so_gör.bind("<Return>", lambda event: self.so_lärt.focus_set())

            self.so_lärt = tk.Entry(self)
            self.so_lärt.grid(row=5, column=2)
            self.so_lärt.bind("<Return>", lambda event: self.so_note.focus_set())

            self.so_note = tk.Entry(self)
            self.so_note.grid(row=5, column=3)
            self.so_note.bind("<Return>", lambda event: self.idh_gör.focus_set())
            
            self.idh_label = tk.Label(self, text="IDH")
            self.idh_label.grid(row=6, column=0, sticky="w")

            self.idh_gör = tk.Entry(self)
            self.idh_gör.grid(row=6, column=1)
            self.idh_gör.bind("<Return>", lambda event: self.idh_lärt.focus_set())

            self.idh_lärt = tk.Entry(self)
            self.idh_lärt.grid(row=6, column=2)
            self.idh_lärt.bind("<Return>", lambda event: self.idh_note.focus_set())

            self.idh_note = tk.Entry(self)
            self.idh_note.grid(row=6, column=3)
            self.idh_note.bind("<Return>", lambda event: self.mu_gör.focus_set())
            
            self.mu_label = tk.Label(self, text="Musik")
            self.mu_label.grid(row=7, column=0, sticky="w")

            self.mu_gör = tk.Entry(self)
            self.mu_gör.grid(row=7, column=1)
            self.mu_gör.bind("<Return>", lambda event: self.mu_lärt.focus_set())

            self.mu_lärt = tk.Entry(self)
            self.mu_lärt.grid(row=7, column=2)
            self.mu_lärt.bind("<Return>", lambda event: self.mu_note.focus_set())

            self.mu_note = tk.Entry(self)
            self.mu_note.grid(row=7, column=3)
            self.mu_note.bind("<Return>", lambda event: self.bi_gör.focus_set())
            
            self.bi_label = tk.Label(self, text="Bild")
            self.bi_label.grid(row=8, column=0, sticky="w")

            self.bi_gör = tk.Entry(self)
            self.bi_gör.grid(row=8, column=1)
            self.bi_gör.bind("<Return>", lambda event: self.bi_lärt.focus_set())

            self.bi_lärt = tk.Entry(self)
            self.bi_lärt.grid(row=8, column=2)
            self.bi_lärt.bind("<Return>", lambda event: self.bi_note.focus_set())

            self.bi_note = tk.Entry(self)
            self.bi_note.grid(row=8, column=3)
            self.bi_note.bind("<Return>", lambda event: self.hkk_gör.focus_set())
            
            self.hkk_label = tk.Label(self, text="Hemkunskap")
            self.hkk_label.grid(row=9, column=0, sticky="w")

            self.hkk_gör = tk.Entry(self)
            self.hkk_gör.grid(row=9, column=1)
            self.hkk_gör.bind("<Return>", lambda event: self.hkk_lärt.focus_set())

            self.hkk_lärt = tk.Entry(self)
            self.hkk_lärt.grid(row=9, column=2)
            self.hkk_lärt.bind("<Return>", lambda event: self.hkk_note.focus_set())
            

            self.hkk_note = tk.Entry(self)
            self.hkk_note.grid(row=9, column=3)
            self.hkk_note.bind("<Return>", lambda event: self.sp_gör.focus_set())
            
            self.sp_label = tk.Label(self, text="Språk")
            self.sp_label.grid(row=10, column=0, sticky="w")

            self.sp_gör = tk.Entry(self)
            self.sp_gör.grid(row=10, column=1)
            self.sp_gör.bind("<Return>", lambda event: self.sp_lärt.focus_set())

            self.sp_lärt = tk.Entry(self)
            self.sp_lärt.grid(row=10, column=2)
            self.sp_lärt.bind("<Return>", lambda event: self.sp_note.focus_set())

            self.sp_note = tk.Entry(self)
            self.sp_note.grid(row=10, column=3)
            self.sp_note.bind("<Return>", lambda event: self.sl_gör.focus_set())
            
            self.sl_label = tk.Label(self, text="Slöjd")
            self.sl_label.grid(row=11, column=0, sticky="w")

            self.sl_gör = tk.Entry(self)
            self.sl_gör.grid(row=11, column=1)
            self.sl_gör.bind("<Return>", lambda event: self.sl_lärt.focus_set())

            self.sl_lärt = tk.Entry(self)
            self.sl_lärt.grid(row=11, column=2)
            self.sl_lärt.bind("<Return>", lambda event: self.sl_note.focus_set())

            self.sl_note = tk.Entry(self)
            self.sl_note.grid(row=11, column=3)
            tk.Entry.pack(self)

        questions(self)
        

        
        #buttons and checkboxes
        self.rest_checked = tk.IntVar()
        self.add_fields_checkbox = tk.Checkbutton(self, text="Har du några rester?", variable=self.rest_checked, command=self.rest)
        self.add_fields_checkbox.grid(row=48, column=0)
        
        self.händelser_checked = tk.IntVar()
        self.add_fields_checkbox = tk.Checkbutton(self, text="Har du några händelser?", command=self.händelser, variable=self.händelser_checked)
        self.add_fields_checkbox.grid(row=49, column=0)
        
        self.mail_message_checked = tk.IntVar()
        self.message = tk.Checkbutton(self, text="Vill du skriva något i mailet?", command=self.mail_message, variable=self.mail_message_checked)
        self.message.grid(row=47, column=0)
        
        self.submit_button = tk.Button(self, text="Submit", command=self.submit)
        self.submit_button.grid(row=50, column=1, columnspan=3, pady=10)
        self.restart_button = tk.Button(self.master, text="Ändra E-epostadress", command=self.restart_program)
        self.restart_button.pack(side=tk.LEFT, anchor=tk.SW, padx=10, pady=10)
        self.restart_button = tk.Button(self.master, text="Vissa senaste hembrevet", command=self.latest_hembrev)
        self.restart_button.pack(side=tk.LEFT, anchor=tk.SW, padx=10, pady=10)
    def latest_hembrev(self):
        

        folder_path = r"C:\Users\liam.suorsa08\hembrev"
        extension = "*.docx"

        # Get a list of all Word documents in the folder
        files = glob.glob(os.path.join(folder_path, extension))

        # Filter out temporary files
        files = [file for file in files if not os.path.basename(file).startswith("~$")]

        # Sort the remaining files by creation time in descending order
        files.sort(key=lambda x: os.path.getctime(x), reverse=True)

        # Open the Word document with the latest creation date
        if files:
            latest_file = files[0]
            os.startfile(latest_file)
        else:
            messagebox.showerror("Error", f"Du har inga hembrev sparade. ")


    def restart_program(self):
        # display a message to the user
        messagebox.showinfo("Ändra E-epostadress", "Starta om programmet för att genomföra ändringar")
        
        # delete the firstime.txt file
        try:
            folder_path = os.path.join(newPath, "hembrev")  # Combine newPath and folder name
            filename = os.path.join(folder_path, "firstime.txt")
            os.remove(filename)
        except FileNotFoundError:
            pass
        
        # exit the program
        self.master.destroy()
        

        

# creating the window

def run(Application):
    root = tk.Tk()
    root.state('zoomed')
    root.title("Hembrev")
    

    #if user presses the X the program confirms it
    def on_closing():
        if messagebox.askokcancel("Stäng", "Är du säker"):
            root.destroy()

    root.protocol("WM_DELETE_WINDOW", on_closing)
    app = Application(master=root)
    app.mainloop()
    
def start(run):
    
    def connect(host='http://google.com'): #checks if connected to internet
        try:
            urllib.request.urlopen(host) 
            return True
        except:
            return False
        
    if connect(): #if connectet to the internet then start the program else error
        filename = os.path.join(fullpath, "firstime.txt")
        if os.path.exists(filename):
            run(Application)
        else:
                setup(start , run)
    else:
        messagebox.showerror("Error", "Du behöver internet för att programmet ska funka")
        logging.info('User didnt have internet')
    
    
    
    
    
    
def setup(start,run):
    folder_path = os.path.join(newPath, "hembrev")  # Combine newPath and folder name
    filename = os.path.join(folder_path, "password-encrypted.txt")
    try:
        os.remove(filename)
    except:
        logging.info('password-encrypted.txt not found')
    folder_path = os.path.join(newPath, "hembrev")  # Combine newPath and folder name
    filename = os.path.join(folder_path, "password.txt")
    try:
        os.remove(filename)
    except:
        logging.info('password.txt not found')
    def submit(start, run):
        def encrypt():
            folder_path = os.path.join(newPath, "hembrev")  # Combine newPath and folder name
            filename = os.path.join(folder_path, file.txt)
            # Read the password from file
            with open(filename, "r") as file:
                password = file.read().strip()

            c = SymmetricCipher(password=password)
            try:
                folder_path = os.path.join(newPath, "hembrev")  # Combine newPath and folder name
                filename = os.path.join(folder_path, "password.txt")
                c.encrypt_file(filename)
                os.remove(filename)
            except OSError as e:
                print(e)
                os.remove("password-encrypted.txt")
                encrypt()
                
        global Aemail
        global Bemail
        global Cemail
        global Demail
        global Eemail
        global login
        global password
        Aemail = email1_var.get()
        Bemail = email2_var.get()
        Cemail = email3_var.get()
        Demail = email4_var.get()
        Eemail = email5_var.get()
        login = login_var.get()
        password = password_var.get()
        config_object = ConfigParser()
        config_object["Emails"] = {
            "email1": Aemail,
            "email2": Bemail,
            "email3": Cemail,
            "email4": Demail,
            "email5": Eemail
        }
        config_object["login"] = {
            "Email": login,
        }
        
        config_object["Settings"] = {
        "auto_message": str(checkbox_var.get()),
    }
        
        if os.path.exists('config.ini'):
            os.remove('config.ini')
            
        def config_write(config_object):
            folder_path = os.path.join(os.path.expanduser('~'), 'AppData', 'Roaming', 'hembrev')  # Combine folder path with the desired folder name
            filename = os.path.join(folder_path, 'config.ini')  # Combine folder path and filename

            with open(filename, 'w') as conf:
                config_object.write(conf)
        
            
        if os.path.exists('password.txt'):
            os.remove('password.txt')
        if not os.path.exists('supergoodfile.txt'):
            f = open("supergoodfile.txt", "x")
            
  
        time.sleep(1)  
        def password_write(password):
            folder_path = os.path.join(newPath, "hembrev")  # Combine newPath and folder name
            filename = os.path.join(folder_path, "password.txt")  # Combine folder path and filename

            with open(filename, 'w') as f:
                f.write(password)
        time.sleep(1)
        try: 
            server = smtplib.SMTP('smtp.office365.com', 587)
            server.connect("smtp.office365.com",587)
            server.ehlo()
            server.starttls()
            server.ehlo()
            server.login(login, password)
            server.quit()
            time.sleep(1)
            config_write(config_object)
            password_write(password)
            encrypt()
            filename = os.path.join(fullpath, "firstime.txt")
            with open(filename, 'w') as file: #create the first time file that checks if the config files is setup
                file.write('This file is only so the program knows if its the first time. DO NOT DELETE')
            root.destroy()
            start(run)
        except smtplib.SMTPAuthenticationError as e:
            print(e)
            messagebox.showerror("Error", "E-postadressen eller lösenordet är felaktigt")
        except Exception as e:
            logging.info(e) 
            messagebox.showerror("Error", f"Något gick fel. Kontakta gärna Liam Suorsa. \n {e}")
            
        
        
        
        
        

    root = tk.Tk()
    root.title("Setup")
    root.geometry("400x300")

    email1_var = tk.StringVar()
    email2_var = tk.StringVar()
    email3_var = tk.StringVar()
    email4_var = tk.StringVar()
    email5_var = tk.StringVar()
    login_var = tk.StringVar()
    password_var = tk.StringVar()
    checkbox_var = tk.IntVar()

    entry = tk.Entry(root, textvariable=email1_var)
    email1_label = tk.Label(root, text="Email 1: ")
    email1_entry = tk.Entry(root, textvariable=email1_var)
    email2_label = tk.Label(root, text="Email 2: ")
    email2_entry = tk.Entry(root, textvariable=email2_var)
    email3_label = tk.Label(root, text="Email 3: ")
    email3_entry = tk.Entry(root, textvariable=email3_var)
    email4_label = tk.Label(root, text="Email 4: ")
    email4_entry = tk.Entry(root, textvariable=email4_var)
    email5_label = tk.Label(root, text="Email 5: ")
    email5_entry = tk.Entry(root, textvariable=email5_var)
    login_label = tk.Label(root, text="Skol e-post: ")
    login_entry = tk.Entry(root, textvariable=login_var)
    password_label = tk.Label(root, text="Lösenord: ")
    password_entry = tk.Entry(root, textvariable=password_var, show="*")
    checkbox_label = tk.Label(root, text="Auto meddelande i mailet: ")
    checkbox = tk.Checkbutton(root, variable=checkbox_var)
    submit_button = tk.Button(root, text="Submit", command=lambda: submit(start, run))
    


    email1_label.grid(row=0, column=0, padx=5, pady=5)
    email1_entry.grid(row=0, column=1, padx=5, pady=5)
    email2_label.grid(row=1, column=0, padx=5, pady=5)
    email2_entry.grid(row=1, column=1, padx=5, pady=5)
    email3_label.grid(row=2, column=0, padx=5, pady=5)
    email3_entry.grid(row=2, column=1, padx=5, pady=5)
    email4_label.grid(row=3, column=0, padx=5, pady=5)
    email4_entry.grid(row=3, column=1, padx=5, pady=5)
    email5_label.grid(row=4, column=0, padx=5, pady=5)
    email5_entry.grid(row=4, column=1, padx=5, pady=5)
    login_label.grid(row=5, column=0, padx=5, pady=5)
    login_entry.grid(row=5, column=1, padx=5, pady=5)
    password_label.grid(row=6, column=0, padx=5, pady=5)
    password_entry.grid(row=6, column=1, padx=5, pady=5)
    checkbox_label.grid(row=7, column=0, padx=5, pady=5)
    checkbox.grid(row=7, column=1, padx=5, pady=5)
    submit_button.grid(row=8, column=1, padx=5, pady=5)

    root.mainloop()
    
    
start(run)


#© Liam Suorsa